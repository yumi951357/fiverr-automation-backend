from fastapi import FastAPI
from pydantic import BaseModel
from typing import Literal, List, Optional
from PIL import Image, ImageDraw, ImageFont
import os, json, io, base64
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from common.audit import append_audit

app = FastAPI(title="Generation Layer", version="1.0.0")

class GigAssetsRequest(BaseModel):
    category: Literal["youtube_script", "data_visualization", "ai_business_plan"]
    tone: Literal["friendly","professional","bold"] = "professional"
    language: Literal["en","zh"] = "en"
    brand_signature: Optional[str] = None

@app.post("/generate_gig_assets")
def generate_gig_assets(req: GigAssetsRequest):
    # simple rules
    if req.category == "youtube_script":
        title = "I will create SEO‑optimized YouTube scripts that increase watch time"
        desc = "Get engaging, SEO‑ready scripts tailored to your niche. Includes keyword‑rich titles, hooks, and CTAs."
        keywords = ["YouTube script", "SEO", "video content", "hooks", "titles"]
    elif req.category == "data_visualization":
        title = "I will design interactive Power BI dashboards for business insights"
        desc = "Clean, interactive dashboards that highlight what matters. Data storytelling, drill‑downs, KPI cards."
        keywords = ["Power BI", "dashboard", "data visualization", "KPI", "insights"]
    else:
        title = "I will craft an AI‑powered business plan ready for investors"
        desc = "Clear market analysis, GTM, and risk mapping. Investor‑ready narrative + optional financial outline."
        keywords = ["business plan", "AI", "investor deck", "market analysis"]

    assets = {"title": title, "description": desc, "keywords": keywords, "language": req.language}
    append_audit("generate_gig_assets", {"category": req.category, "tone": req.tone, "language": req.language})
    return assets

class TemplateRequest(BaseModel):
    service_type: Literal["youtube_script","data_visualization","business_plan"]
    structure: Optional[str] = None

@app.post("/build_template")
def build_template(req: TemplateRequest):
    os.makedirs("storage/templates", exist_ok=True)
    if req.service_type == "youtube_script":
        path = "storage/templates/script_template.docx"
        from docx import Document
        doc = Document()
        doc.add_heading("YouTube Script Template", 0)
        doc.add_paragraph("Title: ")
        doc.add_paragraph("Hook (0–10s): ")
        doc.add_paragraph("Main Points: ")
        doc.add_paragraph("CTA: ")
        doc.add_paragraph("\n\n" + os.environ.get("BRAND_SIGNATURE", ""))
        doc.save(path)
    elif req.service_type == "data_visualization":
        # simple placeholder Excel via CSV; user can import to BI tools
        path = "storage/templates/dashboard_template.csv"
        with open(path, "w", encoding="utf-8") as f:
            f.write("date,revenue,users\n2025-01-01,1000,25\n2025-01-02,1200,30\n")
    else:
        path = "storage/templates/business_plan_template.pdf"
        c = canvas.Canvas(path, pagesize=A4)
        w, h = A4
        c.setFont("Helvetica-Bold", 18)
        c.drawString(72, h-72, "AI-Powered Business Plan Template")
        c.setFont("Helvetica", 12)
        y = h-110
        sections = ["Executive Summary","Problem & Solution","Market & ICP","GTM Strategy","Operations","Financial Outline","Risks"]
        for s in sections:
            c.drawString(72, y, f"- {s}")
            y -= 20
        c.drawString(72, 72, os.environ.get("BRAND_SIGNATURE",""))
        c.showPage(); c.save()

    append_audit("build_template", {"service_type": req.service_type})
    return {"template_path": path}

class CoverRequest(BaseModel):
    label: str
    filename: str

@app.post("/generate_cover")
def generate_cover(req: CoverRequest):
    os.makedirs("storage/covers", exist_ok=True)
    path = f"storage/covers/{req.filename}"
    img = Image.new("RGB", (1080, 608), (20, 24, 36))
    d = ImageDraw.Draw(img)
    # simple headline
    d.rectangle([(40, 440), (1040, 560)], outline=(255,255,255), width=4)
    d.text((60, 60), req.label, fill=(255,255,255))
    brand = os.environ.get("BRAND_SIGNATURE","")
    d.text((60, 520), brand, fill=(200,200,200))
    img.save(path)
    append_audit("generate_cover", {"filename": req.filename})
    return {"cover_path": path}
